from pathlib import Path
import sys

from docx import Document
from docx.shared import Inches


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: fix_word_pagination.py input.docx output.docx")
        return 2

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    doc = Document(str(input_path))

    section = doc.sections[0]
    max_width = section.page_width - section.left_margin - section.right_margin
    max_height = Inches(4.8)

    for paragraph in doc.paragraphs:
        fmt = paragraph.paragraph_format
        fmt.keep_with_next = False
        fmt.keep_together = False
        fmt.page_break_before = False

    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            for cant_split in tr_pr.xpath("./w:cantSplit"):
                tr_pr.remove(cant_split)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fmt = paragraph.paragraph_format
                    fmt.keep_with_next = False
                    fmt.keep_together = False
                    fmt.page_break_before = False

    for shape in doc.inline_shapes:
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = int(shape.width * ratio)
            shape.height = int(shape.height * ratio)
        if shape.height > max_height:
            ratio = max_height / shape.height
            shape.width = int(shape.width * ratio)
            shape.height = int(shape.height * ratio)

    doc.save(str(output_path))
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
