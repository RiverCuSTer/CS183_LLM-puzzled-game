from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
import shutil
import tempfile
import xml.etree.ElementTree as ET

from docx import Document
from docx.shared import Inches


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}


def qn(prefix: str, tag: str) -> str:
    return f"{{{NS[prefix]}}}{tag}"


def remove_intermediate_section_breaks(docx_path: Path) -> None:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        with ZipFile(docx_path) as zin:
            zin.extractall(tmp_path)

        document_xml = tmp_path / "word" / "document.xml"
        tree = ET.parse(document_xml)
        root = tree.getroot()

        body = root.find(qn("w", "body"))
        if body is not None:
            # Keep the final body-level sectPr, remove paragraph-level sectPr that
            # creates unexpected section/page jumps in templates.
            for p in list(body.findall(qn("w", "p"))):
                ppr = p.find(qn("w", "pPr"))
                if ppr is None:
                    continue
                sect = ppr.find(qn("w", "sectPr"))
                if sect is not None:
                    ppr.remove(sect)

        tree.write(document_xml, encoding="utf-8", xml_declaration=True)

        new_path = docx_path.with_suffix(".tmp.docx")
        with ZipFile(new_path, "w", ZIP_DEFLATED) as zout:
            for file_path in tmp_path.rglob("*"):
                if file_path.is_file():
                    zout.write(file_path, file_path.relative_to(tmp_path).as_posix())
        shutil.move(str(new_path), str(docx_path))


def normalize_docx(input_path: Path, output_path: Path) -> None:
    doc = Document(str(input_path))

    for style in doc.styles:
        try:
            fmt = style.paragraph_format
        except Exception:
            continue
        fmt.keep_with_next = False
        fmt.keep_together = False
        fmt.page_break_before = False

    for paragraph in doc.paragraphs:
        fmt = paragraph.paragraph_format
        fmt.keep_with_next = False
        fmt.keep_together = False
        fmt.page_break_before = False
        fmt.widow_control = False

    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            for cant_split in tr_pr.xpath("./w:cantSplit"):
                tr_pr.remove(cant_split)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fmt = paragraph.paragraph_format
                    fmt.keep_with_next = False
                    fmt.keep_together = False
                    fmt.page_break_before = False
                    fmt.widow_control = False

    section = doc.sections[0]
    max_width = section.page_width - section.left_margin - section.right_margin
    max_height = Inches(4.4)
    for shape in doc.inline_shapes:
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = int(shape.width * ratio)
            shape.height = int(shape.height * ratio)
        if shape.height > max_height:
            ratio = max_height / shape.height
            shape.width = int(shape.width * ratio)
            shape.height = int(shape.height * ratio)

    doc.save(str(output_path))
    remove_intermediate_section_breaks(output_path)


if __name__ == "__main__":
    src = Path(r"C:\Users\Lenovo\Desktop\cs183-2026-report-template-word.docx")
    dst = Path(r"C:\Users\Lenovo\Desktop\cs183-2026-report-template-word_fixed_v2.docx")
    normalize_docx(src, dst)
    print(dst)
